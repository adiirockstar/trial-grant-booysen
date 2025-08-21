import streamlit as st
import os
import yaml
from pathlib import Path

def load_config():
    config_path = os.path.join(os.path.dirname(__file__), "../config.yaml")
    with open(config_path, "r") as f:
        return yaml.safe_load(f)

def process_uploaded_file(uploaded_file):
    """Process uploaded file and extract text content"""
    try:
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension in ['txt', 'md']:
            # Text files
            content = str(uploaded_file.read(), "utf-8")
            
        elif file_extension == 'pdf':
            # PDF files - you'll need PyPDF2
            try:
                import PyPDF2
                import io
                
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
                content = ""
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
                    
            except ImportError:
                st.error("PyPDF2 not available. PDF upload disabled.")
                return None
                
        elif file_extension == 'docx':
            # DOCX files - you'll need python-docx
            try:
                from docx import Document
                import io
                
                doc = Document(io.BytesIO(uploaded_file.read()))
                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
            except ImportError:
                st.error("python-docx not available. DOCX upload disabled.")
                return None
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return None
            
        # Create document object
        class UploadedDocument:
            def __init__(self, content, filename):
                self.page_content = content
                self.metadata = {"source": f"uploaded_{filename}", "type": "uploaded"}
        
        return UploadedDocument(content, uploaded_file.name)
        
    except Exception as e:
        st.error(f"Error processing {uploaded_file.name}: {str(e)}")
        return None

def load_frontend():
    cfg = load_config() if os.path.exists("config.yaml") else {}
    
    # Initialize session state for uploaded documents
    if 'uploaded_docs' not in st.session_state:
        st.session_state.uploaded_docs = []
    if 'upload_counter' not in st.session_state:
        st.session_state.upload_counter = 0
    
    # --------- Sidebar: Settings and Upload ---------
    with st.sidebar:
        st.header("⚙️ Settings")
        
        # Mode selection
        mode = st.radio(
            "Response Mode",
            ["Interview", "Story", "FastFacts", "HumbleBrag", "Reflect"],
            index=0
        )
        st.caption("Mode tweaks tone & structure of answers.")
        
        st.divider()
        
        # Document upload section
        st.header("📁 Add Training Data")
        st.caption("Upload documents to expand Grant's knowledge base")
        st.info("💡 Files are processed in-memory for this session")
        
        # File uploader
        uploaded_files = st.file_uploader(
            "Choose files",
            type=['txt', 'pdf', 'docx', 'md'],
            accept_multiple_files=True,
            help="Supported formats: TXT, PDF, DOCX, Markdown",
            key=f"file_uploader_{st.session_state.upload_counter}"
        )
        
        # Process uploaded files
        if uploaded_files:
            new_docs = []
            
            for uploaded_file in uploaded_files:
                # Check if already processed
                already_uploaded = any(
                    doc.metadata.get("source") == f"uploaded_{uploaded_file.name}" 
                    for doc in st.session_state.uploaded_docs
                )
                
                if not already_uploaded:
                    doc = process_uploaded_file(uploaded_file)
                    if doc:
                        new_docs.append(doc)
                        st.success(f"✅ Processed: {uploaded_file.name}")
                else:
                    st.info(f"📄 Already uploaded: {uploaded_file.name}")
            
            if new_docs:
                st.session_state.uploaded_docs.extend(new_docs)
                # Increment counter to force re-initialization
                if 'docs_changed' not in st.session_state:
                    st.session_state.docs_changed = 0
                st.session_state.docs_changed += 1
                
                st.success(f"🎉 Added {len(new_docs)} new documents!")
        
        # Show document statistics
        total_original = len([f for f in os.listdir("data") if f.lower().endswith(('.txt', '.pdf', '.docx', '.md'))]) if os.path.exists("data") else 0
        total_uploaded = len(st.session_state.uploaded_docs)
        total_docs = total_original + total_uploaded
        
        st.metric("📊 Total Documents", total_docs, delta=total_uploaded if total_uploaded > 0 else None)
        
        if total_uploaded > 0:
            st.write("**Uploaded this session:**")
            for doc in st.session_state.uploaded_docs:
                filename = doc.metadata["source"].replace("uploaded_", "")
                content_length = len(doc.page_content)
                st.write(f"• {filename} ({content_length:,} chars)")
        
        st.divider()
        
        # Session management
        st.header("🔄 Session Management")
        
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ Clear Uploads"):
                st.session_state.uploaded_docs = []
                st.session_state.docs_changed = st.session_state.get('docs_changed', 0) + 1
                st.success("Cleared uploaded documents!")
                st.rerun()
        
        with col2:
            if st.button("🔄 Refresh"):
                st.session_state.docs_changed = st.session_state.get('docs_changed', 0) + 1
                st.success("Refreshing knowledge base!")
                st.rerun()
        
        # Download processed content (for backup)
        if st.session_state.uploaded_docs:
            if st.button("💾 Download Session Data"):
                combined_content = "\n\n---\n\n".join([
                    f"FILE: {doc.metadata['source']}\n\n{doc.page_content}"
                    for doc in st.session_state.uploaded_docs
                ])
                
                st.download_button(
                    label="📄 Download All Uploaded Content",
                    data=combined_content,
                    file_name="uploaded_session_data.txt",
                    mime="text/plain"
                )
                
        with st.expander("🔗 Fun links about me to ask me more about"):
            res = cfg.get("resources", {})
            if res:
                if res.get("spotify_playlist"):
                    st.markdown(f"- **Worship Spotify playlist:** [{res['spotify_playlist']}]({res['spotify_playlist']})")
                if res.get("chess_profile"):
                    label = res.get("chess_username", res["chess_username"])
                    st.markdown(f"- **Chess.com:** [{label}]({res['chess_profile']}) Make me a friend request ;)")
                if res.get("karate_achievements"):
                    st.markdown(f"- **Karate achievements:** [{res['karate_achievements']}]({res['karate_achievements']})")
            else:
                st.info("Add links in `config.yaml → resources` to show them here.")


    # Return mode and uploaded documents
    return mode, st.session_state.uploaded_docs